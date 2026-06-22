import requests
from tqdm import tqdm
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading

# ── Config ──────────────────────────────────────────────────────────────────
BASE_URL   = "https://work.brookfieldproperties.pro/public/OrderImages/"
START      = 1779480000
END        = 1779499999         # adjust if you want a smaller test window
EXTENSIONS = ["", ".jfif"]
TIMEOUT    = 0.3                # VPS optimized
MAX_WORKERS = 400               # increased to 400
OUTPUT_DOC = os.path.join(os.path.dirname(__file__), "found_images.docx")
# ────────────────────────────────────────────────────────────────────────────

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
}

# Create a session with connection pooling for better performance
session = requests.Session()
session.headers.update(HEADERS)
adapter = requests.adapters.HTTPAdapter(
    pool_connections=400,
    pool_maxsize=400,
    max_retries=0
)
session.mount('http://', adapter)
session.mount('https://', adapter)

valid_urls = []
valid_urls_lock = threading.Lock()
total = END - START + 1

print(f"Brookfield Properties — Image URL Scraper")
print(f"Range : {START:,} → {END:,}  ({total:,} numbers)")
print(f"Workers: {MAX_WORKERS}")
print(f"Output: {OUTPUT_DOC}\n")

def check_number(num):
    """Check a single number with all extensions"""
    for ext in EXTENSIONS:
        url = f"{BASE_URL}{num}{ext}"
        try:
            r = session.head(url, allow_redirects=False, timeout=TIMEOUT)
            if r.status_code == 200:
                with valid_urls_lock:
                    valid_urls.append(url)
                return url
        except:
            pass
    return None

with tqdm(
    total=total,
    unit="num",
    unit_scale=True,
    dynamic_ncols=True,
    bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}, {rate_fmt}] Found: {postfix}",
) as pbar:
    pbar.set_postfix_str("0")
    
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = {executor.submit(check_number, num): num for num in range(START, END + 1)}
        
        for future in as_completed(futures):
            result = future.result()
            if result:
                tqdm.write(f"[FOUND] {result}")
                pbar.set_postfix_str(str(len(valid_urls)))
            pbar.update(1)

# ── Build Word document ──────────────────────────────────────────────────────
doc = Document()

# Title
title = doc.add_heading("Brookfield Properties — Found Image URLs", level=1)
title.runs[0].font.size = Pt(16)

# Meta info
doc.add_paragraph(
    f"Scan completed : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    f"Range scanned  : {START:,} – {END:,}\n"
    f"Total found    : {len(valid_urls)}"
)
doc.add_paragraph("")

if valid_urls:
    doc.add_heading("Found URLs", level=2)
    for i, url in enumerate(valid_urls, 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(url)
        run.font.size = Pt(9)
        # Make each URL a clickable hyperlink inside the paragraph text
        run.font.color.rgb = None   # keep default; Word auto-detects http links
else:
    doc.add_paragraph("No valid image URLs were found in the scanned range.")

doc.save(OUTPUT_DOC)

print(f"\n{'─'*50}")
print(f"Scan complete. Found {len(valid_urls)} valid image URL(s).")
print(f"Results saved to: {OUTPUT_DOC}")
